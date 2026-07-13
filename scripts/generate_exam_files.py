"""
K12英语试卷多格式导出脚本
依赖安装：
pip install weasyprint markdown python-docx
"""
import markdown
import os
from weasyprint import HTML, CSS
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def split_exam_content(full_markdown: str):
    """拆分试卷正文、答题卡、答案解析、分析报告"""
    parts = {}
    sections = full_markdown.split("## ")
    current_part = "exam"
    exam_content = []
    answer_sheet = []
    analysis = []
    report = []

    for section in sections:
        if section.startswith("答题卡"):
            current_part = "answer_sheet"
            continue
        elif section.startswith("答案与解析"):
            current_part = "analysis"
            continue
        elif section.startswith("试卷分析报告"):
            current_part = "report"
            continue
        
        if current_part == "exam":
            exam_content.append("## " + section)
        elif current_part == "answer_sheet":
            answer_sheet.append("## " + section)
        elif current_part == "analysis":
            analysis.append("## " + section)
        elif current_part == "report":
            report.append("## " + section)
    
    return {
        "exam": "\n".join(exam_content).lstrip("# "),
        "answer_sheet": "\n".join(answer_sheet),
        "analysis": "\n".join(analysis),
        "report": "\n".join(report)
    }

def get_print_css():
    """标准打印样式CSS"""
    return CSS(string="""
        @page {
            size: A4;
            margin: 2cm 1.8cm;
            @top-center {
                font-size: 12px;
                color: #333;
            }
            @bottom-center {
                content: "第 " counter(page) " 页 / 共 " counter(pages) " 页";
                font-size: 10px;
                color: #666;
            }
        }
        body {
            font-family: "SimSun", "宋体", serif;
            font-size: 14px;
            line-height: 1.6;
        }
        h1 {
            text-align: center;
            font-size: 22px;
            margin-bottom: 10px;
            font-weight: bold;
        }
        h2 {
            font-size: 16px;
            margin-top: 18px;
            border-bottom: 1px solid #333;
            padding-bottom: 4px;
        }
        h3 {
            font-size: 14px;
            margin-top: 12px;
        }
        .options {
            margin-left: 20px;
        }
        .answer-blank {
            border-bottom: 1px solid #000;
            display: inline-block;
            width: 120px;
            margin: 0 5px;
        }
        table {
            width: 100%;
            border-collapse: collapse;
            margin: 10px 0;
        }
        td, th {
            border: 1px solid #333;
            padding: 6px;
            text-align: center;
        }
    """)

def generate_pdfs(full_markdown: str, output_dir: str = "./output"):
    """生成全套PDF文件"""
    os.makedirs(output_dir, exist_ok=True)
    parts = split_exam_content(full_markdown)
    print_css = get_print_css()

    def md_to_pdf(md_content, filename):
        html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
        HTML(string=html).write_pdf(
            os.path.join(output_dir, filename),
            stylesheets=[print_css]
        )
    
    md_to_pdf(parts["exam"], "试卷正文.pdf")
    md_to_pdf("## 答题卡\n" + parts["answer_sheet"], "答题卡.pdf")
    md_to_pdf("## 答案与解析\n" + parts["analysis"], "答案解析.pdf")
    if parts["report"].strip():
        md_to_pdf("## 试卷分析报告\n" + parts["report"], "试卷分析报告.pdf")
    
    print(f"PDF生成完成，已保存至 {output_dir} 目录")

def generate_word(full_markdown: str, output_dir: str = "./output"):
    """生成可编辑Word文档"""
    os.makedirs(output_dir, exist_ok=True)
    doc = Document()
    
    # 页面设置
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    # 写入内容（简化版Markdown转Word）
    lines = full_markdown.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.line_spacing = 1.6
            for run in p.runs:
                run.font.name = "宋体"
                run.font.size = Pt(14)
    
    doc.save(os.path.join(output_dir, "试卷正文.docx"))
    print("Word文档生成完成")

if __name__ == "__main__":
    with open("exam.md", "r", encoding="utf-8") as f:
        content = f.read()
    generate_pdfs(content)
    generate_word(content)
